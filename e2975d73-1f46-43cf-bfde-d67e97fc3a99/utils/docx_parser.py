import os
import re
from typing import Dict, Any, List
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from config.settings import DOWNLOAD_DIR
from utils.logger import logger, log_error_with_context


class DocxParser:
    def __init__(self):
        self.download_dir = DOWNLOAD_DIR

    def parse(self, file_path: str) -> Dict[str, Any]:
        result = {
            'success': False,
            'text': '',
            'tables': [],
            'metadata': {},
            'paragraphs': [],
            'error': None
        }

        if not os.path.exists(file_path):
            result['error'] = f"File not found: {file_path}"
            logger.error(result['error'])
            return result

        if not file_path.lower().endswith('.docx'):
            result['error'] = f"Not a .docx file: {file_path}"
            logger.warning(result['error'])
            return self._parse_legacy_doc(file_path, result)

        try:
            doc = Document(file_path)

            result['metadata'] = self._extract_metadata(doc)

            paragraphs_text, paragraphs_data = self._extract_paragraphs(doc)
            result['paragraphs'] = paragraphs_data

            tables_data = self._extract_tables(doc)
            result['tables'] = tables_data

            full_text = paragraphs_text
            if tables_data:
                full_text += '\n\n' + '\n'.join([
                    self._table_to_text(table) for table in tables_data
                ])

            result['text'] = self._clean_text(full_text)
            result['success'] = True

            logger.info(f"Successfully parsed DOCX: {file_path}, paragraphs: {len(paragraphs_data)}, tables: {len(tables_data)}")

        except PackageNotFoundError as e:
            log_error_with_context(logger, e, f"Invalid DOCX package: {file_path}")
            result['error'] = f"Invalid document: {str(e)}"
            return self._parse_legacy_doc(file_path, result)
        except Exception as e:
            log_error_with_context(logger, e, f"Failed to parse DOCX: {file_path}")
            result['error'] = str(e)

        return result

    def _parse_legacy_doc(self, file_path: str, result: Dict[str, Any]) -> Dict[str, Any]:
        if file_path.lower().endswith('.doc'):
            logger.info(f"Attempting to parse legacy .doc file: {file_path}")
            try:
                import subprocess
                result['text'] = self._extract_text_with_catdoc(file_path)
                if result['text'].strip():
                    result['success'] = True
                    result['error'] = None
                    logger.info(f"Successfully parsed .doc with catdoc: {file_path}")
            except Exception as e:
                log_error_with_context(logger, e, f"Failed to parse .doc file: {file_path}")
        return result

    def _extract_text_with_catdoc(self, file_path: str) -> str:
        try:
            import subprocess
            result = subprocess.run(
                ['catdoc', '-d', 'utf-8', file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return self._clean_text(result.stdout)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        try:
            import subprocess
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return self._clean_text(result.stdout)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        return ''

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        metadata = {}
        try:
            core_props = doc.core_properties
            props = [
                'author', 'category', 'comments', 'content_status',
                'created', 'identifier', 'keywords', 'language',
                'last_modified_by', 'last_printed', 'modified',
                'revision', 'subject', 'title', 'version'
            ]
            for prop in props:
                value = getattr(core_props, prop, None)
                if value is not None:
                    if hasattr(value, 'strftime'):
                        metadata[prop] = value.strftime('%Y-%m-%d %H:%M:%S')
                    else:
                        metadata[prop] = str(value)
        except Exception as e:
            log_error_with_context(logger, e, "Failed to extract DOCX metadata")

        return metadata

    def _extract_paragraphs(self, doc: Document) -> tuple:
        paragraphs_data = []
        full_text_parts = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ''
                run_styles = []
                for run in para.runs:
                    if run.bold:
                        run_styles.append('bold')
                    if run.italic:
                        run_styles.append('italic')
                    if run.underline:
                        run_styles.append('underline')

                para_data = {
                    'index': i,
                    'text': text,
                    'style': style,
                    'is_heading': 'Heading' in style if style else False,
                    'heading_level': int(style.replace('Heading ', '')) if style and style.startswith('Heading ') else None,
                    'styles': list(set(run_styles))
                }
                paragraphs_data.append(para_data)
                full_text_parts.append(text)

        return '\n'.join(full_text_parts), paragraphs_data

    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        tables_data = []

        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'index': table_idx,
                'rows': []
            }

            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    row_data.append({
                        'row': row_idx,
                        'col': cell_idx,
                        'text': cell_text
                    })
                table_data['rows'].append(row_data)

            if any(any(cell['text'] for cell in row) for row in table_data['rows']):
                tables_data.append(table_data)

        return tables_data

    def _table_to_text(self, table_data: Dict[str, Any]) -> str:
        text_parts = []
        for row in table_data['rows']:
            row_text = ' | '.join([cell['text'] for cell in row])
            text_parts.append(row_text)
        return '\n'.join(text_parts)

    def _clean_text(self, text: str) -> str:
        if not text:
            return ''

        text = text.replace('\x00', '')
        text = text.replace('\u3000', ' ')
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)

        text = re.sub(
            r'[^\x00-\x7F\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\n\r\t ]+',
            ' ', text
        )

        return text.strip()

    def extract_title(self, paragraphs: List[Dict[str, Any]], fallback: str = '') -> str:
        if not paragraphs:
            return fallback

        for para in paragraphs:
            if para.get('is_heading') and para.get('heading_level') == 1:
                return para['text']

        for para in paragraphs:
            if para.get('is_heading'):
                return para['text']

        for para in paragraphs:
            text = para['text']
            if 5 < len(text) < 100 and 'bold' in para.get('styles', []):
                return text

        for para in paragraphs[:10]:
            text = para['text']
            if 5 < len(text) < 100:
                if not re.match(r'^[0-9\s\.\-]+$', text):
                    return text

        return fallback

    def extract_summary(self, text: str, max_length: int = 300) -> str:
        if not text:
            return ''

        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]

        for para in paragraphs:
            if any(keyword in para for keyword in ['为贯彻', '为落实', '根据', '按照', '现就', '为加强', '为规范']):
                if len(para) <= max_length:
                    return para
                return para[:max_length] + '...'

        for para in paragraphs:
            if len(para) > 50:
                return para[:max_length] + ('...' if len(para) > max_length else '')

        return text[:max_length]


docx_parser = DocxParser()


def parse_docx(file_path: str) -> Dict[str, Any]:
    return docx_parser.parse(file_path)

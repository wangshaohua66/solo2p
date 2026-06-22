"""
docx_parser.py
================================================================================
Word 文档（.docx）解析模块：使用 python-docx 解析 Word 文档内容，
支持表格提取、段落文本解析，以及从 Word 格式的人员名单、工资单、
缴费通知单等文档中提取结构化数据。

特性：
  1. 解析 docx 文档中的所有段落文本
  2. 提取 docx 文档中的所有表格，自动识别表头行
  3. 支持 Word 表格跨行/跨列合并单元格
  4. 统一返回类似 Excel 解析器的结构化数据（行列表）
  5. 与 config.yaml 的 field_mapping 保持一致的字段映射
  6. 支持从缴费通知单等非表格文本中提取关键字段（金额、日期等）
"""

from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict, List, Optional, Sequence, Tuple

try:
    from docx import Document
    from docx.table import Table, _Cell
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None
    Table = None
    _Cell = None

from excel_parser import ExcelParser

logger = logging.getLogger(__name__)


@dataclass
class DocxParsed:
    """docx 解析结果。"""
    file_path: str
    paragraphs: List[str] = field(default_factory=list)
    tables: List[List[List[Any]]] = field(default_factory=list)
    structured_rows: List[Dict[str, Any]] = field(default_factory=list)
    table_type: str = "unknown"  # "wage" / "personnel" / "notice" / "unknown"
    header_row: Optional[int] = None
    column_map: Dict[str, str] = field(default_factory=dict)
    raw_columns: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    @property
    def row_count(self) -> int:
        return len(self.structured_rows)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "file_path": self.file_path,
            "paragraphs": self.paragraphs,
            "structured_rows": self.structured_rows,
            "table_type": self.table_type,
            "header_row": self.header_row,
            "column_map": self.column_map,
            "row_count": self.row_count,
        }


@dataclass
class NoticeInfo:
    """从缴费通知单提取的关键字段。"""
    file_path: str
    amount: Optional[float] = None
    pay_date: Optional[str] = None
    due_date: Optional[str] = None
    unit_name: Optional[str] = None
    insurance_types: List[str] = field(default_factory=list)
    raw_text: str = ""
    parse_ok: bool = False


class DocxParser:
    """docx 文档解析器。"""

    def __init__(self, config: Dict[str, Any]) -> None:
        if not HAS_DOCX:
            raise ImportError("缺少 python-docx 依赖，请运行：pip install python-docx")
        self.config = config
        # 复用 ExcelParser 的字段映射与表头检测逻辑
        self._excel_parser = ExcelParser(config)

    # ---------------------------- 公共入口 ----------------------------

    def parse_file(self, file_path: str,
                    table_hint: Optional[str] = None) -> DocxParsed:
        """解析 docx 文件，自动识别是否为表格类文档还是通知单类文档。

        Args:
            file_path: .docx 文件路径
            table_hint: "wage" / "personnel" / "notice" 类型提示
        """
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"docx 文件不存在：{file_path}")
        ext = os.path.splitext(file_path)[1].lower()
        if ext != ".docx":
            raise ValueError(f"不支持的文件类型：{ext}，仅支持 .docx")

        doc = Document(file_path)
        result = DocxParsed(file_path=file_path)

        # 1. 提取所有段落
        result.paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(result.paragraphs)
        logger.info("docx 段落数：%d", len(result.paragraphs))

        # 2. 提取所有表格
        for table_idx, table in enumerate(doc.tables):
            table_data = self._extract_table(table)
            result.tables.append(table_data)
            logger.debug("表格 %d：%d 行 × %d 列", table_idx, len(table_data),
                        len(table_data[0]) if table_data else 0)

        # 3. 自动选择最合适的表格进行结构化解析
        if result.tables:
            # 选择含最多行的表格作为数据表
            best_table = max(result.tables, key=len)
            self._parse_structured(best_table, result, table_hint)

        # 4. 若未识别到表格或为通知单，尝试段落文本解析
        if result.table_type == "unknown" or table_hint == "notice":
            notice = self._parse_notice(full_text)
            notice.file_path = file_path
            if notice.parse_ok:
                result.table_type = "notice"
                result.warnings.append(
                    f"识别为缴费通知单：金额={notice.amount}，日期={notice.pay_date}")

        logger.info("docx 解析完成：类型=%s，行数=%d",
                    result.table_type, result.row_count)
        return result

    def parse_notice(self, file_path: str) -> NoticeInfo:
        """便捷方法：仅解析缴费通知单（从段落文本提取关键字段）。"""
        doc = Document(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        notice = self._parse_notice(full_text)
        notice.file_path = file_path
        return notice

    def parse_wage_table(self, file_path: str) -> DocxParsed:
        """便捷方法：仅解析工资表（取第一张识别为 wage 的表）。"""
        result = self.parse_file(file_path, table_hint="wage")
        if result.table_type != "wage":
            logger.warning("未在 %s 中识别到工资表，返回空数据", file_path)
        return result

    def parse_personnel_table(self, file_path: str) -> DocxParsed:
        """便捷方法：仅解析人员变动表。"""
        result = self.parse_file(file_path, table_hint="personnel")
        if result.table_type != "personnel":
            logger.warning("未在 %s 中识别到人员变动表，返回空数据", file_path)
        return result

    # ---------------------------- 表格提取 ----------------------------

    def _extract_table(self, table: Table) -> List[List[Any]]:
        """将 docx Table 转为二维列表，填充合并单元格。"""
        # 先读取原始数据
        raw = []
        for row in table.rows:
            raw.append([cell.text for cell in row.cells])
        if not raw:
            return []
        # 处理合并单元格：python-docx 自动展开合并单元格，
        # 但我们需要确保行列对齐
        max_cols = max(len(r) for r in raw)
        for row in raw:
            while len(row) < max_cols:
                row.append("")
        return raw

    def _parse_structured(self, table_data: List[List[Any]],
                          result: DocxParsed,
                          table_hint: Optional[str]) -> None:
        """将二维表格数据复用 ExcelParser 的表头检测与字段映射逻辑。"""
        import pandas as pd
        df = pd.DataFrame(table_data)
        # 使用 _parse_dataframe 方法（注意这是 ExcelParser 内部方法，我们直接复制关键逻辑）
        df = df.fillna("").astype(str)
        header_row, table_type, column_map = self._detect_header(df, table_hint)
        result.header_row = header_row
        result.table_type = table_type
        result.column_map = column_map
        if header_row is None or not column_map:
            return
        header_values = [str(v) for v in df.iloc[header_row].tolist()]
        result.raw_columns = header_values
        # 取表头以下数据行
        body = df.iloc[header_row + 1:].reset_index(drop=True)
        body.columns = header_values
        for _, row in body.iterrows():
            record = self._excel_parser._row_to_record(row, column_map, table_type)  # noqa
            if record and self._excel_parser._has_any_value(record):  # noqa
                result.structured_rows.append(record)

    def _detect_header(self, df, hint: Optional[str]) -> Tuple[Optional[int], str, Dict[str, str]]:
        # 直接委托给 ExcelParser 内部的检测方法，因为它是私有方法
        return self._excel_parser._detect_header(df, hint)

    # ---------------------------- 通知单解析 ----------------------------

    def _parse_notice(self, full_text: str) -> NoticeInfo:
        """从段落文本中提取缴费通知单关键字段。"""
        notice = NoticeInfo(file_path="")
        notice.raw_text = full_text
        ocr_cfg = self.config.get("ocr", {})
        amount_pattern = ocr_cfg.get(
            "amount_pattern", r'(?:大写|金额)[:：]?\s*([0-9,]+\.\d{2})')
        date_pattern = ocr_cfg.get(
            "date_pattern", r'(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?)')

        # 金额匹配多个金额匹配
        # 金额
        amount_matches = re.findall(amount_pattern, full_text)
        for raw in amount_matches:
            cleaned = raw.replace(",", "")
            try:
                notice.amount = float(cleaned)
                break
            except ValueError:
                continue
        if notice.amount is None:
            all_amounts = re.findall(r'(\d{1,3}(?:,\d{3})*\.\d{2})', full_text)
            for raw in all_amounts or []:
                try:
                    notice.amount = float(raw.replace(",", ""))
                    break
                except ValueError:
                    continue

        # 日期
        date_matches = re.findall(date_pattern, full_text)
        for raw in date_matches:
            normalized = self._normalize_date(raw)
            if normalized:
                notice.pay_date = normalized
                break

        # 截止日期
        due_match = re.search(r'(?:缴费截止|缴款期限|缴费期限)[:：\s]*([^\n]+)', full_text)
        if due_match:
            notice.due_date = self._normalize_date(due_match.group(1))

        # 单位名称
        unit_match = re.search(r'(?:参保单位|缴费单位|单位名称)[:：\s]*([^\n]+)', full_text)
        if unit_match:
            notice.unit_name = unit_match.group(1).strip()

        # 费种匹配
        from pdf_parser import _INSURANCE_KEYWORDS  # 复用 pdf_parser 的关键词
        for ins_type, keywords in _INSURANCE_KEYWORDS.items():
            if any(kw in full_text for kw in keywords):
                if ins_type not in notice.insurance_types:
                    notice.insurance_types.append(ins_type)

        notice.parse_ok = notice.amount is not None or notice.pay_date
        return notice

    # ---------------------------- 工具方法 ----------------------------

    @staticmethod
    def _normalize_date(text: str) -> Optional[str]:
        if not text:
            return None
        text = str(text).strip()
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y年%m月%d日", "%Y%m%d"):
            try:
                return datetime.strptime(text, fmt).strftime("%Y-%m-%d")
            except ValueError:
                continue
        nums = re.findall(r"\d+", text)
        if len(nums) >= 3:
            try:
                y, mo, d = int(nums[0]), int(nums[1]), int(nums[2])
                if 1900 <= y <= 2100 and 1 <= mo <= 12 and 1 <= d <= 31:
                    return f"{y:04d}-{mo:02d}-{d:02d}"
            except (ValueError, IndexError):
                pass
        return None


def to_records(parsed: DocxParsed) -> List[Dict[str, Any]]:
    """便捷函数：将 DocxParsed 转为纯字典列表。"""
    return list(parsed.structured_rows)

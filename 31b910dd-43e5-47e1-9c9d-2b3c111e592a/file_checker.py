import re
import io
import time
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from config import DRUG_TYPE_CONFIG, SIGNATURE_PAGES, OVERVIEW_FIELDS, TEMPLATE_DIR
from logger import logger

try:
    import pdfplumber
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PDF库未安装，PDF相关功能将降级为文件存在性检查")

try:
    import cv2
    import numpy as np
    from PIL import Image
    CV_AVAILABLE = True
except ImportError:
    CV_AVAILABLE = False
    logger.warning("OpenCV/Pillow未安装，签字盖章识别功能将跳过")

try:
    import pyautogui
    PYAUTOGUI_AVAILABLE = True
except ImportError:
    PYAUTOGUI_AVAILABLE = False
    logger.warning("PyAutoGUI未安装，GUI截图功能将降级")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx未安装，Word文档解析功能将跳过")

SEAL_TEMPLATE_DIR = TEMPLATE_DIR / "seal_templates"
SIG_TEMPLATE_DIR = TEMPLATE_DIR / "sig_templates"
SEAL_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
SIG_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

SEAL_TEMPLATE_THRESHOLD = 0.65
SIG_TEMPLATE_THRESHOLD = 0.55
SEAL_BLUR_VARIANCE_THRESHOLD = 80.0
SIGNATURE_CONSISTENCY_THRESHOLD = 0.70

ROMAN_NUMERAL_MAP = {
    "I": 1, "II": 2, "III": 3, "IV": 4, "V": 5,
    "VI": 6, "VII": 7, "VIII": 8, "IX": 9, "X": 10,
    "XI": 11, "XII": 12, "XIII": 13, "XIV": 14, "XV": 15,
    "XVI": 16, "XVII": 17, "XVIII": 18, "XIX": 19, "XX": 20,
    "i": 1, "ii": 2, "iii": 3, "iv": 4, "v": 5,
    "vi": 6, "vii": 7, "viii": 8, "ix": 9, "x": 10,
}


@dataclass
class FileIssue:
    issue_type: str
    module: str
    severity: str
    description: str
    file_path: str = ""
    page: Optional[int] = None
    suggestion: str = ""
    is_common: bool = False
    matched_issue_id: Optional[int] = None


@dataclass
class PageInfo:
    file_path: str
    module: str
    page_index: int
    page_number: Optional[str]
    is_roman: bool = False


@dataclass
class OverviewInfo:
    drug_name: str = ""
    application_type: str = ""
    specification: str = ""
    applicant: str = ""
    fields_found: Dict[str, bool] = field(default_factory=dict)


def _load_templates(template_dir: Path) -> List[np.ndarray]:
    templates: List[np.ndarray] = []
    if not CV_AVAILABLE:
        return templates
    for ext in ("*.png", "*.jpg", "*.jpeg", "*.bmp"):
        for p in template_dir.glob(ext):
            try:
                img = cv2.imread(str(p), cv2.IMREAD_COLOR)
                if img is not None:
                    templates.append(img)
                    logger.debug(f"加载模板: {p.name} ({img.shape})")
            except Exception as e:
                logger.warning(f"加载模板失败 {p}: {e}")
    return templates


def _generate_default_seal_templates() -> None:
    if not CV_AVAILABLE:
        return
    sizes = [(120, 120), (150, 150), (180, 180)]
    for idx, (w, h) in enumerate(sizes):
        img = np.ones((h, w, 3), dtype=np.uint8) * 255
        center = (w // 2, h // 2)
        radius = min(w, h) // 2 - 5
        cv2.circle(img, center, radius, (0, 0, 200), 3)
        cv2.circle(img, center, radius - 8, (0, 0, 200), 1)
        inner_r = radius - 15
        if inner_r > 10:
            cv2.circle(img, center, inner_r, (0, 0, 180), 1)
        cv2.ellipse(img, center, (radius - 12, radius - 12), 0, -30, 210, (0, 0, 180), 1)
        path = SEAL_TEMPLATE_DIR / f"default_seal_{idx}.png"
        if not path.exists():
            cv2.imwrite(str(path), img)
            logger.debug(f"生成默认印章模板: {path.name}")


def _generate_default_sig_templates() -> None:
    if not CV_AVAILABLE:
        return
    for idx, (w, h) in enumerate([(180, 60), (160, 50), (200, 70)]):
        img = np.ones((h, w, 3), dtype=np.uint8) * 255
        y_mid = h // 2
        cv2.line(img, (10, y_mid), (w - 10, y_mid), (0, 0, 0), 1)
        cv2.ellipse(img, (w // 3, y_mid - 5), (30, 15), -10, 0, 360, (30, 30, 30), 2)
        cv2.ellipse(img, (2 * w // 3, y_mid + 3), (25, 12), 10, 0, 360, (30, 30, 30), 2)
        path = SIG_TEMPLATE_DIR / f"default_sig_{idx}.png"
        if not path.exists():
            cv2.imwrite(str(path), img)
            logger.debug(f"生成默认签字模板: {path.name}")


_generate_default_seal_templates()
_generate_default_sig_templates()


class FileChecker:
    def __init__(self, files: List[Path], drug_type: str = "chemical",
                 files_by_module: Optional[Dict[str, List[Path]]] = None) -> None:
        self.files = files
        self.drug_type = drug_type
        self.drug_config = DRUG_TYPE_CONFIG.get(drug_type, DRUG_TYPE_CONFIG["chemical"])
        self.files_by_module = files_by_module or {}
        self.issues: List[FileIssue] = []
        self.page_records: List[PageInfo] = []
        self.overview_info: Optional[OverviewInfo] = None
        self._seal_templates: Optional[List[np.ndarray]] = None
        self._sig_templates: Optional[List[np.ndarray]] = None
        self._signature_features: Dict[str, np.ndarray] = {}

    def _roman_to_int(self, s: str) -> Optional[int]:
        return ROMAN_NUMERAL_MAP.get(s)

    def _get_seal_templates(self) -> List[np.ndarray]:
        if self._seal_templates is None:
            self._seal_templates = _load_templates(SEAL_TEMPLATE_DIR)
        return self._seal_templates

    def _get_sig_templates(self) -> List[np.ndarray]:
        if self._sig_templates is None:
            self._sig_templates = _load_templates(SIG_TEMPLATE_DIR)
        return self._sig_templates

    def check_file_naming(self) -> List[FileIssue]:
        logger.info("开始文件命名规范校验")
        pattern = re.compile(self.drug_config["naming_pattern"])
        naming_issues: List[FileIssue] = []

        for file_path in self.files:
            filename = file_path.name
            if not pattern.match(filename):
                naming_issues.append(FileIssue(
                    issue_type="INVALID_FILENAME",
                    module="file_naming",
                    severity="DEFECT",
                    description=f"文件名不符合CTD命名规范: {filename}",
                    file_path=str(file_path),
                    suggestion=f"请按{self.drug_config['label']}格式修正，"
                               f"例如: {self.drug_config['module_prefix']}1-001-资料名称.pdf",
                ))
            else:
                match = pattern.match(filename)
                if match:
                    module_num = match.group(1)
                    if module_num and int(module_num) > 5:
                        naming_issues.append(FileIssue(
                            issue_type="INVALID_MODULE_CODE",
                            module="file_naming",
                            severity="DEFECT",
                            description=f"模块编号超出CTD范围(1-5): {filename}",
                            file_path=str(file_path),
                            suggestion="CTD模块编号应为1到5",
                        ))

        self.issues.extend(naming_issues)
        logger.info(f"命名校验完成，发现 {len(naming_issues)} 个问题")
        return naming_issues

    def _extract_page_number(self, text: str) -> Tuple[Optional[str], bool]:
        page_patterns = [
            r"第\s*(\d+|[IVXivx]+)\s*页",
            r"第\s*(\d+|[IVXivx]+)\s*/\s*\d+\s*页",
            r"Page\s+(\d+)",
            r"Page\s+(\d+)\s+of\s+\d+",
            r"[^\d](\d{1,4})[^\d]\s*$",
            r"^\s*(\d+)\s*$",
            r"^\s*([IVXivx]{1,5})\s*$",
        ]

        for pattern in page_patterns:
            matches = re.findall(pattern, text)
            if matches:
                page_str = matches[-1] if isinstance(matches[-1], str) else matches[-1][0]
                is_roman = bool(re.match(r"^[IVXivx]+$", page_str))
                return page_str, is_roman
        return None, False

    def _extract_page_number_from_pdf(self, pdf_path: Path) -> List[PageInfo]:
        if not PDF_AVAILABLE:
            return []

        module = self._infer_module_from_path(pdf_path)
        pages: List[PageInfo] = []

        try:
            with pdfplumber.open(str(pdf_path)) as pdf:
                total = len(pdf.pages)
                sample_indices = list(range(min(5, total))) + list(range(max(0, total - 5), total))
                if total > 20:
                    sample_indices += list(range(5, total - 5, max(1, total // 20)))

                for idx in sorted(set(sample_indices)):
                    if idx >= total:
                        continue
                    try:
                        page = pdf.pages[idx]
                        text = page.extract_text() or ""
                        footer_region = ""
                        if page.height:
                            footer_bbox = (0, page.height * 0.85, page.width, page.height)
                            try:
                                footer = page.within_bbox(footer_bbox).extract_text()
                                footer_region = footer or ""
                            except Exception:
                                footer_region = text[-200:] if len(text) > 200 else text
                        else:
                            footer_region = text

                        search_text = (footer_region + "\n" + text[-300:]) if len(text) > 300 else text
                        page_num, is_roman = self._extract_page_number(search_text)

                        pages.append(PageInfo(
                            file_path=str(pdf_path),
                            module=module,
                            page_index=idx,
                            page_number=page_num,
                            is_roman=is_roman,
                        ))
                    except Exception as e:
                        logger.warning(f"提取页码失败 {pdf_path} 第{idx}页: {e}")
                        pages.append(PageInfo(
                            file_path=str(pdf_path),
                            module=module,
                            page_index=idx,
                            page_number=None,
                        ))
        except Exception as e:
            logger.error(f"PDF解析失败 {pdf_path}: {e}", exception=e)
            self.issues.append(FileIssue(
                issue_type="PDF_PARSE_ERROR",
                module="page_continuity",
                severity="DEFECT",
                description=f"PDF文件解析失败，降级为存在性检查: {pdf_path.name}",
                file_path=str(pdf_path),
                suggestion="请检查PDF文件是否损坏或加密",
            ))
        return pages

    def _infer_module_from_path(self, path: Path) -> str:
        parts = path.parts
        for part in parts:
            match = re.match(r"Module(\d)", part, re.IGNORECASE)
            if match:
                return f"Module{match.group(1)}"
        return "Unknown"

    def _page_number_to_int(self, page_info: PageInfo) -> Optional[int]:
        if page_info.page_number is None:
            return None
        if page_info.is_roman:
            return self._roman_to_int(page_info.page_number)
        try:
            return int(page_info.page_number)
        except ValueError:
            return None

    def check_page_continuity(self) -> List[FileIssue]:
        logger.info("开始页码连续性验证")
        continuity_issues: List[FileIssue] = []
        all_pages: List[PageInfo] = []

        pdf_files = [f for f in self.files if f.suffix.lower() == ".pdf"]
        for pdf_file in pdf_files:
            pages = self._extract_page_number_from_pdf(pdf_file)
            all_pages.extend(pages)
            self.page_records.extend(pages)

        modules_pages: Dict[str, List[PageInfo]] = {}
        for p in all_pages:
            modules_pages.setdefault(p.module, []).append(p)

        for module, pages in modules_pages.items():
            pages_sorted = sorted(pages, key=lambda x: (x.file_path, x.page_index))
            numbered_pages = [(p, self._page_number_to_int(p)) for p in pages_sorted]
            numbered_pages = [(p, n) for p, n in numbered_pages if n is not None]

            roman_pages = [p for p, n in numbered_pages if p.is_roman]
            arabic_pages = [p for p, n in numbered_pages if not p.is_roman]

            for page_group, group_name in [(roman_pages, "罗马数字"), (arabic_pages, "阿拉伯数字")]:
                seen: Dict[int, PageInfo] = {}
                last_num: Optional[int] = None
                last_page: Optional[PageInfo] = None

                for p in page_group:
                    num = self._page_number_to_int(p)
                    if num is None:
                        continue

                    if num in seen:
                        prev = seen[num]
                        continuity_issues.append(FileIssue(
                            issue_type="DUPLICATE_PAGE_NUMBER",
                            module="page_continuity",
                            severity="DEFECT",
                            description=f"{module} 内页码重号: {p.page_number}({group_name}), "
                                        f"分别出现在 {Path(prev.file_path).name} 第{prev.page_index + 1}页 和 "
                                        f"{Path(p.file_path).name} 第{p.page_index + 1}页",
                            file_path=p.file_path,
                            page=p.page_index + 1,
                            suggestion="请修正重复页码，确保每页页码唯一",
                        ))
                    seen[num] = p

                    if last_num is not None:
                        if num < last_num:
                            continuity_issues.append(FileIssue(
                                issue_type="PAGE_ORDER_ERROR",
                                module="page_continuity",
                                severity="DEFECT",
                                description=f"{module} 内页码顺序错误: 前一页{last_num}({group_name}) 后一页{num}",
                                file_path=p.file_path,
                                page=p.page_index + 1,
                                suggestion="请检查文件排序与页码标注",
                            ))
                        elif num > last_num + 1:
                            missing = list(range(last_num + 1, num))
                            continuity_issues.append(FileIssue(
                                issue_type="MISSING_PAGE_NUMBER",
                                module="page_continuity",
                                severity="DEFECT",
                                description=f"{module} 内页码缺号: 缺失页码 {missing}({group_name}), "
                                            f"在 {Path(last_page.file_path if last_page else p).name} 之后",
                                file_path=p.file_path,
                                page=p.page_index + 1,
                                suggestion=f"请补充缺失的页码或修正跳号问题",
                            ))
                    last_num = num
                    last_page = p

        total_checked = len(all_pages)
        matched_pages = sum(1 for p in all_pages if p.page_number is not None)
        if total_checked > 0:
            accuracy = matched_pages / total_checked
            logger.info(f"页码识别覆盖率: {accuracy:.2%} ({matched_pages}/{total_checked})")
            if accuracy < 0.80:
                continuity_issues.append(FileIssue(
                    issue_type="LOW_PAGE_RECOGNITION",
                    module="page_continuity",
                    severity="SUGGESTION",
                    description=f"页码识别覆盖率较低: {accuracy:.0%}，建议人工复核",
                    suggestion="建议优化PDF页眉页脚格式，使用标准页码标注",
                ))

        self.issues.extend(continuity_issues)
        logger.info(f"页码验证完成，发现 {len(continuity_issues)} 个问题")
        return continuity_issues

    def _render_pdf_page_to_image(self, pdf_path: Path, page_idx: int) -> Optional[np.ndarray]:
        if not CV_AVAILABLE:
            return None

        if PDF_AVAILABLE:
            try:
                with pdfplumber.open(str(pdf_path)) as pdf:
                    if page_idx >= len(pdf.pages):
                        return None
                    page = pdf.pages[page_idx]
                    pil_img = page.to_image(resolution=200)
                    img_array = np.array(pil_img.original)
                    if len(img_array.shape) == 2:
                        img_array = cv2.cvtColor(img_array, cv2.COLOR_GRAY2BGR)
                    elif img_array.shape[2] == 4:
                        img_array = cv2.cvtColor(img_array, cv2.COLOR_RGBA2BGR)
                    elif img_array.shape[2] == 3:
                        img_array = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                    return img_array
            except Exception as e:
                logger.debug(f"pdfplumber渲染页面失败，尝试PyPDF2: {pdf_path} page {page_idx}: {e}")

        if PDF_AVAILABLE:
            try:
                reader = PdfReader(str(pdf_path))
                if page_idx >= len(reader.pages):
                    return None
                page = reader.pages[page_idx]
                xObject = page.get("/Resources", {}).get("/XObject", {})
                for obj_name in xObject:
                    obj = xObject[obj_name].get_object()
                    if obj.get("/Subtype") == "/Image":
                        width = obj.get("/Width", 0)
                        height = obj.get("/Height", 0)
                        color_space = obj.get("/ColorSpace", "/DeviceRGB")
                        data = obj.get_data()
                        if width > 0 and height > 0 and len(data) > 0:
                            if "/DeviceRGB" in str(color_space):
                                img = np.frombuffer(data, dtype=np.uint8).reshape((height, width, 3))
                                return cv2.cvtColor(img, cv2.COLOR_RGB2BGR)
                            elif "/DeviceGray" in str(color_space):
                                img = np.frombuffer(data, dtype=np.uint8).reshape((height, width))
                                return cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)
            except Exception as e:
                logger.debug(f"PyPDF2图像提取失败: {pdf_path} page {page_idx}: {e}")

        logger.debug(f"无法渲染PDF页面为图像，跳过: {pdf_path} page {page_idx}")
        return None

    def _screenshot_pdf_page(self, pdf_path: Path, page_idx: int) -> Optional[np.ndarray]:
        if not PYAUTOGUI_AVAILABLE or not CV_AVAILABLE:
            return None

        try:
            import subprocess
            subprocess.Popen(
                ["open", str(pdf_path)],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            )
            time.sleep(2.0)

            screen_w, screen_h = pyautogui.size()
            screenshot = pyautogui.screenshot()
            img = cv2.cvtColor(np.array(screenshot), cv2.COLOR_RGB2BGR)

            subprocess.Popen(
                ["osascript", "-e", 'tell application "Preview" to close front window'],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            )
            return img
        except Exception as e:
            logger.warning(f"PyAutoGUI截图失败 {pdf_path} page {page_idx}: {e}")
            return None

    def _acquire_page_image(self, pdf_path: Path, page_idx: int) -> Optional[np.ndarray]:
        img = self._render_pdf_page_to_image(pdf_path, page_idx)
        if img is not None:
            return img

        if page_idx in (0, 1):
            img = self._screenshot_pdf_page(pdf_path, page_idx)
            if img is not None:
                return img

        return None

    def _match_template(
        self, image: np.ndarray, template: np.ndarray, threshold: float
    ) -> List[Tuple[int, int, float]]:
        if not CV_AVAILABLE:
            return []

        th, tw = template.shape[:2]
        ih, iw = image.shape[:2]
        if th > ih or tw > iw:
            resized = template.copy()
            scale = min(iw * 0.8 / tw, ih * 0.8 / th)
            new_w = max(1, int(tw * scale))
            new_h = max(1, int(th * scale))
            resized = cv2.resize(resized, (new_w, new_h), interpolation=cv2.INTER_AREA)
            if new_h > ih or new_w > iw:
                return []
            template = resized
            th, tw = new_h, new_w

        img_gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        tmpl_gray = cv2.cvtColor(template, cv2.COLOR_BGR2GRAY)

        result = cv2.matchTemplate(img_gray, tmpl_gray, cv2.TM_CCOEFF_NORMED)
        _, max_val, _, max_loc = cv2.minMaxLoc(result)

        matches: List[Tuple[int, int, float]] = []
        if max_val >= threshold:
            matches.append((max_loc[0], max_loc[1], float(max_val)))

        locs = np.where(result >= threshold * 0.95)
        for pt in zip(*locs[::-1]):
            score = float(result[pt[1], pt[0]])
            is_duplicate = False
            for mx, my, ms in matches:
                if abs(pt[0] - mx) < tw // 2 and abs(pt[1] - my) < th // 2:
                    is_duplicate = True
                    break
            if not is_duplicate:
                matches.append((int(pt[0]), int(pt[1]), score))

        return matches

    def _detect_seal_by_template(self, image: np.ndarray) -> List[Tuple[int, int, int, int, float]]:
        if not CV_AVAILABLE:
            return []

        hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)
        lower_red1 = np.array([0, 43, 46])
        upper_red1 = np.array([10, 255, 255])
        lower_red2 = np.array([156, 43, 46])
        upper_red2 = np.array([180, 255, 255])
        mask1 = cv2.inRange(hsv, lower_red1, upper_red1)
        mask2 = cv2.inRange(hsv, lower_red2, upper_red2)
        red_mask = cv2.bitwise_or(mask1, mask2)
        kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
        red_mask = cv2.morphologyEx(red_mask, cv2.MORPH_CLOSE, kernel)
        red_image = cv2.bitwise_and(image, image, mask=red_mask)

        templates = self._get_seal_templates()
        all_matches: List[Tuple[int, int, int, int, float]] = []

        for tmpl in templates:
            matches = self._match_template(red_image, tmpl, SEAL_TEMPLATE_THRESHOLD)
            for mx, my, score in matches:
                th, tw = tmpl.shape[:2]
                all_matches.append((mx, my, tw, th, score))

        contours, _ = cv2.findContours(red_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        for cnt in contours:
            x, y, cw, ch = cv2.boundingRect(cnt)
            area = cv2.contourArea(cnt)
            aspect_ratio = cw / max(ch, 1)
            if 60 < cw < 300 and 60 < ch < 300 and 2000 < area < 80000 and 0.5 < aspect_ratio < 2.0:
                circularity = 4 * np.pi * area / (cw * ch * 4) if cw * ch > 0 else 0
                if circularity > 0.3:
                    overlap = False
                    for mx, my, mw, mh, ms in all_matches:
                        if (abs(x - mx) < max(cw, mw) * 0.5 and abs(y - my) < max(ch, mh) * 0.5):
                            overlap = True
                            break
                    if not overlap:
                        all_matches.append((x, y, cw, ch, 0.5))

        return all_matches

    def _detect_signature_by_template(self, image: np.ndarray) -> List[Tuple[int, int, int, int, float]]:
        if not CV_AVAILABLE:
            return []

        h, w = image.shape[:2]
        search_area = image[int(h * 0.55):h, int(w * 0.2):w]

        gray = cv2.cvtColor(search_area, cv2.COLOR_BGR2GRAY)
        _, binary = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)
        contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        contour_regions: List[Tuple[int, int, int, int]] = []
        for cnt in contours:
            x, y, cw, ch = cv2.boundingRect(cnt)
            area = cv2.contourArea(cnt)
            if 30 < cw < 400 and 15 < ch < 150 and 400 < area < 20000:
                contour_regions.append((x + int(w * 0.2), y + int(h * 0.55), cw, ch))

        templates = self._get_sig_templates()
        all_matches: List[Tuple[int, int, int, int, float]] = []

        for tmpl in templates:
            matches = self._match_template(search_area, tmpl, SIG_TEMPLATE_THRESHOLD)
            for mx, my, score in matches:
                th, tw = tmpl.shape[:2]
                all_matches.append((mx + int(w * 0.2), my + int(h * 0.55), tw, th, score))

        for cx, cy, cw, ch in contour_regions:
            overlap = False
            for mx, my, mw, mh, ms in all_matches:
                if (abs(cx - mx) < max(cw, mw) * 0.5 and abs(cy - my) < max(ch, mh) * 0.5):
                    overlap = True
                    break
            if not overlap:
                all_matches.append((cx, cy, cw, ch, 0.4))

        return all_matches

    def _compute_blur_variance(self, image_region: np.ndarray) -> float:
        if not CV_AVAILABLE:
            return 0.0
        gray = cv2.cvtColor(image_region, cv2.COLOR_BGR2GRAY) if len(image_region.shape) == 3 else image_region
        return float(cv2.Laplacian(gray, cv2.CV_64F).var())

    def _check_seal_clarity(self, image: np.ndarray, seal_regions: List[Tuple[int, int, int, int, float]]) -> List[FileIssue]:
        clarity_issues: List[FileIssue] = []
        for x, y, w, h, score in seal_regions:
            if w < 20 or h < 20:
                continue
            try:
                region = image[y:y + h, x:x + w]
                blur_var = self._compute_blur_variance(region)
                if blur_var < SEAL_BLUR_VARIANCE_THRESHOLD:
                    clarity_issues.append(FileIssue(
                        issue_type="FUZZY_SEAL",
                        module="signature_seal",
                        severity="DEFECT",
                        description=f"印章疑似模糊（清晰度指标 {blur_var:.1f}，阈值 {SEAL_BLUR_VARIANCE_THRESHOLD:.0f}）"
                                    f"，位置({x},{y}) 大小{w}x{h}",
                        suggestion="请检查印章是否清晰可辨，建议重新盖章并扫描",
                    ))
            except Exception as e:
                logger.debug(f"印章清晰度检测异常: {e}")
        return clarity_issues

    def _compute_signature_features(self, image: np.ndarray, region: Tuple[int, int, int, int]) -> Optional[np.ndarray]:
        if not CV_AVAILABLE:
            return None
        x, y, w, h = region
        if w < 10 or h < 10:
            return None
        try:
            sig_area = image[y:y + h, x:x + w]
            gray = cv2.cvtColor(sig_area, cv2.COLOR_BGR2GRAY)
            resized = cv2.resize(gray, (64, 64), interpolation=cv2.INTER_AREA)
            _, binary = cv2.threshold(resized, 180, 255, cv2.THRESH_BINARY_INV)
            moments = cv2.moments(binary)
            if moments["m00"] == 0:
                return None
            hu_moments = cv2.HuMoments(moments).flatten()
            log_hu = -np.sign(hu_moments) * np.log10(np.abs(hu_moments) + 1e-10)
            return log_hu
        except Exception:
            return None

    def _check_signature_consistency(self, file_path: str, page_idx: int,
                                      image: np.ndarray,
                                      sig_regions: List[Tuple[int, int, int, int, float]]) -> List[FileIssue]:
        consistency_issues: List[FileIssue] = []
        key = f"{file_path}_p{page_idx}"

        if len(sig_regions) == 0:
            return consistency_issues

        for i, (x, y, w, h, score) in enumerate(sig_regions):
            features = self._compute_signature_features(image, (x, y, w, h))
            if features is None:
                continue
            self._signature_features[f"{key}_s{i}"] = features

        if len(self._signature_features) < 2:
            return consistency_issues

        feature_list = list(self._signature_features.items())
        for i in range(len(feature_list)):
            for j in range(i + 1, len(feature_list)):
                key_i, feat_i = feature_list[i]
                key_j, feat_j = feature_list[j]
                if feat_i is None or feat_j is None:
                    continue
                diff = np.linalg.norm(feat_i - feat_j)
                max_diff = max(np.linalg.norm(feat_i), np.linalg.norm(feat_j), 1e-10)
                similarity = 1.0 - (diff / max_diff)
                if similarity < SIGNATURE_CONSISTENCY_THRESHOLD:
                    consistency_issues.append(FileIssue(
                        issue_type="SIGNATURE_INCONSISTENCY",
                        module="signature_seal",
                        severity="DEFECT",
                        description=f"签名一致性异常（相似度 {similarity:.2%}，阈值 {SIGNATURE_CONSISTENCY_THRESHOLD:.0%}），"
                                    f"可能存在代签问题",
                        file_path=file_path,
                        page=page_idx + 1,
                        suggestion="请人工核实签字是否为同一人签署，确认无代签情况",
                    ))

        return consistency_issues

    def check_signature_seal(self) -> List[FileIssue]:
        logger.info("开始签字盖章智能识别（模板匹配 + 图像分析）")
        sig_issues: List[FileIssue] = []
        if not CV_AVAILABLE:
            logger.warning("OpenCV不可用，跳过签字盖章识别")
            return sig_issues

        pdf_files = [f for f in self.files if f.suffix.lower() == ".pdf"]
        checked = 0
        sig_count = 0
        seal_count = 0

        for pdf_path in pdf_files[:20]:
            total_pages = 0
            try:
                if PDF_AVAILABLE:
                    reader = PdfReader(str(pdf_path))
                    total_pages = len(reader.pages)
                else:
                    total_pages = 1
            except Exception:
                continue

            check_indices: set = set()
            for cp in SIGNATURE_PAGES["cover_pages"]:
                if cp - 1 < total_pages:
                    check_indices.add(cp - 1)
            for tp in SIGNATURE_PAGES["toc_pages"]:
                if tp - 1 < total_pages:
                    check_indices.add(tp - 1)
            end_count = max(1, int(total_pages * SIGNATURE_PAGES["end_pages_ratio"]))
            for i in range(total_pages - end_count, total_pages):
                if i >= 0:
                    check_indices.add(i)

            for page_idx in sorted(check_indices):
                checked += 1
                img = self._acquire_page_image(pdf_path, page_idx)
                if img is None:
                    logger.debug(f"无法获取页面图像: {pdf_path.name} page {page_idx}")
                    continue

                seal_regions = self._detect_seal_by_template(img)
                sig_regions = self._detect_signature_by_template(img)
                seal_count += len(seal_regions)
                sig_count += len(sig_regions)

                if page_idx == 0 and not seal_regions:
                    sig_issues.append(FileIssue(
                        issue_type="MISSING_SEAL",
                        module="signature_seal",
                        severity="FATAL",
                        description=f"{pdf_path.name} 首页未检测到单位公章（模板匹配+颜色检测均未发现）",
                        file_path=str(pdf_path),
                        page=1,
                        suggestion="首页必须加盖申请单位鲜章，请补充",
                    ))

                if page_idx == 0 and not sig_regions and not seal_regions:
                    sig_issues.append(FileIssue(
                        issue_type="MISSING_SIGNATURE",
                        module="signature_seal",
                        severity="DEFECT",
                        description=f"{pdf_path.name} 首页未检测到法定代表人签字或授权签字（模板匹配未命中）",
                        file_path=str(pdf_path),
                        page=1,
                        suggestion="首页应有法定代表人或授权人签字",
                    ))

                fuzzy_issues = self._check_seal_clarity(img, seal_regions)
                sig_issues.extend(fuzzy_issues)

                consistency_issues = self._check_signature_consistency(
                    str(pdf_path), page_idx, img, sig_regions
                )
                sig_issues.extend(consistency_issues)

        logger.info(f"签字盖章检查完成: 检查{checked}页，检测到签字{sig_count}处，印章{seal_count}处")
        logger.info(f"发现 {len(sig_issues)} 个签字盖章问题")
        self.issues.extend(sig_issues)
        return sig_issues

    def _parse_docx_fields(self, docx_path: Path) -> OverviewInfo:
        info = OverviewInfo()
        if not DOCX_AVAILABLE:
            return info

        field_patterns = {
            "drug_name": [r"药品名称[::\s]+([^\n\r]+)", r"通用名称[::\s]+([^\n\r]+)"],
            "application_type": [r"申报类型[::\s]+([^\n\r]+)", r"注册分类[::\s]+([^\n\r]+)"],
            "specification": [r"规格[::\s]+([^\n\r]+)", r"剂型规格[::\s]+([^\n\r]+)"],
            "applicant": [r"申请人[::\s]+([^\n\r]+)", r"申请机构[::\s]+([^\n\r]+)"],
        }

        try:
            doc = Document(str(docx_path))
            full_text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    cells_text = "\t".join(cell.text for cell in row.cells)
                    full_text += "\n" + cells_text

            for field, patterns in field_patterns.items():
                info.fields_found[field] = False
                for pattern in patterns:
                    match = re.search(pattern, full_text)
                    if match:
                        value = match.group(1).strip().strip(":：").strip()
                        if value:
                            setattr(info, field, value)
                            info.fields_found[field] = True
                            break
        except Exception as e:
            logger.error(f"解析Word文档失败 {docx_path}: {e}", exception=e)
        return info

    def check_overview_content(self) -> Tuple[List[FileIssue], OverviewInfo]:
        logger.info("开始综述文档内容抽检")
        overview_issues: List[FileIssue] = []
        info = OverviewInfo()

        docx_files = [f for f in self.files if f.suffix.lower() in (".doc", ".docx")]
        overview_docs = [f for f in docx_files if any(k in f.name for k in
                         ["综述", "overview", "2.3", "2-3", "目录", "申请"])]

        if not overview_docs and docx_files:
            overview_docs = docx_files[:3]

        for docx_path in overview_docs:
            parsed = self._parse_docx_fields(docx_path)
            if any(parsed.fields_found.values()):
                info = parsed
                break

        missing_fields = [f for f in OVERVIEW_FIELDS if not info.fields_found.get(f, False)]
        if missing_fields:
            field_names = {
                "drug_name": "药品名称",
                "application_type": "申报类型",
                "specification": "规格",
                "applicant": "申请人",
            }
            missing_cn = [field_names.get(f, f) for f in missing_fields]
            overview_issues.append(FileIssue(
                issue_type="MISSING_OVERVIEW_FIELD",
                module="overview_check",
                severity="DEFECT",
                description=f"综述资料未检测到关键字段: {', '.join(missing_cn)}",
                file_path=str(overview_docs[0]) if overview_docs else "",
                suggestion="请在综述资料中补充药品名称、申报类型、规格、申请人等信息",
            ))

        self.overview_info = info
        self.issues.extend(overview_issues)
        logger.info(f"综述抽检完成，发现 {len(overview_issues)} 个问题")
        return overview_issues, info

    def cross_validate_with_directory(self, ctd_info: Dict[str, Any]) -> List[FileIssue]:
        logger.info("开始目录与综述交叉校验")
        cross_issues: List[FileIssue] = []
        info = self.overview_info or OverviewInfo()

        if info.drug_name and ctd_info.get("project_name"):
            project_name = str(ctd_info["project_name"])
            if info.drug_name not in project_name and project_name not in info.drug_name:
                cross_issues.append(FileIssue(
                    issue_type="NAME_MISMATCH",
                    module="cross_validate",
                    severity="DEFECT",
                    description=f"综述中药品名称'{info.drug_name}'与项目名称'{project_name}'不一致",
                    suggestion="请确认药品名称并保持申报资料中名称统一",
                ))

        self.issues.extend(cross_issues)
        logger.info(f"交叉校验完成，发现 {len(cross_issues)} 个问题")
        return cross_issues

    def run_all_checks(self, project_context: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        start_time = time.time()
        project_context = project_context or {}

        self.check_file_naming()
        self.check_page_continuity()
        self.check_signature_seal()
        self.check_overview_content()
        if project_context:
            self.cross_validate_with_directory(project_context)

        result = {
            "total_issues": len(self.issues),
            "by_module": {},
            "by_severity": {},
            "overview_info": self.overview_info,
            "duration": time.time() - start_time,
        }

        for issue in self.issues:
            result["by_module"][issue.module] = result["by_module"].get(issue.module, 0) + 1
            result["by_severity"][issue.severity] = result["by_severity"].get(issue.severity, 0) + 1

        return result
